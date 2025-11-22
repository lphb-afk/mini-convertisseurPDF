from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image, ImageDraw
from docx import Document
import os

# Créer un PDF simple
def create_sample_pdf():
    c = canvas.Canvas("sample.pdf", pagesize=letter)
    c.drawString(100, 750, "Ceci est un document PDF de test.")
    c.drawString(100, 730, "Il contient du texte simple pour les tests.")
    c.drawString(100, 710, "Utilisez-le pour tester les conversions.")
    c.save()
    print("sample.pdf créé")

# Créer un document Word simple
def create_sample_docx():
    doc = Document()
    doc.add_heading('Document Word de Test', 0)
    doc.add_paragraph('Ceci est un paragraphe de test.')
    doc.add_paragraph('Il peut être utilisé pour tester la conversion Word vers PDF.')
    doc.save('sample.docx')
    print("sample.docx créé")

# Créer une image simple
def create_sample_image():
    img = Image.new('RGB', (400, 300), color='lightblue')
    draw = ImageDraw.Draw(img)
    draw.text((50, 50), "Image de Test", fill='black')
    draw.text((50, 100), "Pour conversion vers PDF", fill='black')
    img.save('sample.png')
    print("sample.png créé")

if __name__ == "__main__":
    create_sample_pdf()
    create_sample_docx()
    create_sample_image()
    print("Fichiers de test créés : sample.pdf, sample.docx, sample.png")